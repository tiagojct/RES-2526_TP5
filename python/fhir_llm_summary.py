#!/usr/bin/env python3
"""
fhir_llm_summary.py

Script de demonstracao: obter dados FHIR de um doente e gerar um resumo
clinico com um LLM (Gemini Flash 2.0 via OpenRouter).

Uso:
    export OPENROUTER_API_KEY="sk-or-..."
    python fhir_llm_summary.py <PATIENT_ID>

Exemplo:
    python fhir_llm_summary.py 131365376

O script:
  1. Obtem todos os recursos do doente via $everything
  2. Extrai e formata Patient, Conditions e Observations
  3. Envia ao LLM para gerar um resumo clinico em portugues
  4. Exporta para JSON, Markdown e Word (.docx)
"""

import sys
import os
import json
import datetime
import requests

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ──────────────────────────────────────────────
# Configuracao
# ──────────────────────────────────────────────

FHIR_BASE = "https://hapi.fhir.org/baseR4"
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "nvidia/nemotron-3-super-120b-a12b:free"

def get_api_key():
    key = os.environ.get("OPENROUTER_API_KEY", "")
    if not key:
        print("ERRO: variavel de ambiente OPENROUTER_API_KEY nao definida.")
        print()
        print("Configura a chave antes de executar:")
        print('  export OPENROUTER_API_KEY="sk-or-..."')
        sys.exit(1)
    return key


# ──────────────────────────────────────────────
# 1. Obter dados do servidor FHIR
# ──────────────────────────────────────────────

def fetch_everything(patient_id):
    """Obtem todos os recursos associados ao doente via $everything."""
    url = f"{FHIR_BASE}/Patient/{patient_id}/$everything"
    print(f"A obter dados de {url} ...")
    resp = requests.get(url, headers={"Accept": "application/fhir+json"}, timeout=30)
    resp.raise_for_status()
    bundle = resp.json()
    entries = bundle.get("entry", [])
    print(f"  Recebidos {len(entries)} recursos.")
    return entries


def extract_resources(entries):
    """Separa os recursos por tipo."""
    patient = None
    conditions = []
    observations = []
    other = []

    for entry in entries:
        resource = entry.get("resource", {})
        rtype = resource.get("resourceType", "")
        if rtype == "Patient":
            patient = resource
        elif rtype == "Condition":
            conditions.append(resource)
        elif rtype == "Observation":
            observations.append(resource)
        else:
            other.append(resource)

    return patient, conditions, observations, other


# ──────────────────────────────────────────────
# 2. Formatar os dados para o prompt
# ──────────────────────────────────────────────

def format_patient(patient):
    if not patient:
        return "Sem dados do doente."
    name_parts = patient.get("name", [{}])[0]
    given = " ".join(name_parts.get("given", []))
    family = name_parts.get("family", "")
    name = f"{given} {family}".strip() or "Desconhecido"
    gender = patient.get("gender", "desconhecido")
    birth = patient.get("birthDate", "desconhecida")
    return f"Nome: {name}\nSexo: {gender}\nData de nascimento: {birth}"


def format_condition(cond):
    coding = cond.get("code", {}).get("coding", [{}])[0]
    display = coding.get("display", "sem descricao")
    code = coding.get("code", "?")
    system = coding.get("system", "")
    status_coding = cond.get("clinicalStatus", {}).get("coding", [{}])[0]
    status = status_coding.get("code", "desconhecido")

    system_label = system.split("/")[-1] if system else "?"
    return f"- {display} (codigo: {code}, sistema: {system_label}, estado: {status})"


def format_observation(obs):
    coding = obs.get("code", {}).get("coding", [{}])[0]
    display = coding.get("display", "sem descricao")
    code = coding.get("code", "?")
    date = obs.get("effectiveDateTime", obs.get("issued", "data desconhecida"))
    status = obs.get("status", "?")

    vq = obs.get("valueQuantity", {})
    if vq:
        value = f"{vq.get('value', '?')} {vq.get('unit', '')}"
    elif obs.get("valueString"):
        value = obs["valueString"]
    elif obs.get("valueCodeableConcept", {}).get("text"):
        value = obs["valueCodeableConcept"]["text"]
    else:
        value = "sem valor registado"

    return f"- [{date}] {display} (LOINC: {code}): {value} (status: {status})"


def build_clinical_context(patient, conditions, observations):
    """Constroi uma representacao textual dos dados clinicos."""
    lines = []
    lines.append("DADOS DO DOENTE:")
    lines.append(format_patient(patient))
    lines.append("")

    lines.append(f"DIAGNOSTICOS ({len(conditions)}):")
    if conditions:
        for c in conditions:
            lines.append(format_condition(c))
    else:
        lines.append("- Nenhum diagnostico registado.")
    lines.append("")

    lines.append(f"OBSERVACOES / RESULTADOS ({len(observations)}):")
    if observations:
        sorted_obs = sorted(
            observations,
            key=lambda o: o.get("effectiveDateTime", o.get("issued", "")),
            reverse=True
        )
        # Limitar a 30 observacoes para nao exceder o contexto
        for o in sorted_obs[:30]:
            lines.append(format_observation(o))
        if len(sorted_obs) > 30:
            lines.append(f"  ... e mais {len(sorted_obs) - 30} observacoes omitidas.")
    else:
        lines.append("- Nenhuma observacao registada.")

    return "\n".join(lines)


# ──────────────────────────────────────────────
# 3. Chamar o LLM via OpenRouter
# ──────────────────────────────────────────────

SYSTEM_PROMPT = """Es um assistente clinico. Recebes dados estruturados FHIR de um doente 
e produzes um resumo clinico conciso em portugues europeu. 

Regras:
- Usa apenas a informacao fornecida. Nao inventes dados.
- Se a informacao for insuficiente, indica-o explicitamente.
- Organiza o resumo em: identificacao do doente, problemas ativos, 
  resultados laboratoriais relevantes, e uma breve apreciacao global.
- Nao uses emojis.
- Nao repitas a mesma informacao varias vezes.
- Se existirem varios resultados do mesmo exame, menciona a tendencia (a subir, estavel, a descer)."""


def call_llm(clinical_context, api_key):
    """Envia os dados ao LLM e devolve o resumo."""
    user_message = (
        "Com base nos seguintes dados FHIR do doente, gera um resumo clinico:\n\n"
        + clinical_context
    )

    print("A enviar dados ao LLM (Gemini Flash 2.0 via OpenRouter)...")
    resp = requests.post(
        OPENROUTER_URL,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": MODEL,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_message},
            ],
            "max_tokens": 1500,
            "temperature": 0.3,
        },
        timeout=60,
    )
    if resp.status_code != 200:
        print(f"ERRO do OpenRouter: {resp.status_code}")
        try:
            print(json.dumps(resp.json(), indent=2, ensure_ascii=False))
        except Exception:
            print(resp.text[:500])
        sys.exit(1)
    data = resp.json()

    # Extrair o texto da resposta
    choices = data.get("choices", [])
    if not choices:
        return "ERRO: o LLM nao devolveu resposta."
    return choices[0].get("message", {}).get("content", "Sem conteudo na resposta.")


# ──────────────────────────────────────────────
# 4. Exportar JSON
# ──────────────────────────────────────────────

def export_json(patient_id, patient, conditions, observations, summary):
    """Exporta os dados originais e o resumo num ficheiro JSON."""
    output = {
        "metadata": {
            "generated_at": datetime.datetime.now().isoformat(),
            "fhir_server": FHIR_BASE,
            "patient_id": patient_id,
            "model": MODEL,
        },
        "fhir_data": {
            "patient": patient,
            "conditions": conditions,
            "observations": observations,
        },
        "clinical_summary": summary,
    }

    filename = f"resumo_clinico_{patient_id}.json"
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"  JSON: {filename}")
    return filename


def export_markdown(patient_id, patient, conditions, observations, summary):
    """Exporta o resumo clínico como ficheiro Markdown."""
    name_parts = patient.get("name", [{}])[0]
    given = " ".join(name_parts.get("given", []))
    family = name_parts.get("family", "")
    name = f"{given} {family}".strip() or "Desconhecido"
    birth = patient.get("birthDate", "desconhecida")
    gender = patient.get("gender", "desconhecido")

    lines = []
    lines.append(f"# Resumo Clínico: {name}")
    lines.append("")
    lines.append(f"Data de geração: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append(f"Modelo: {MODEL}")
    lines.append(f"Servidor FHIR: {FHIR_BASE}")
    lines.append("")
    lines.append("## Identificação do doente")
    lines.append("")
    lines.append(f"- Nome: {name}")
    lines.append(f"- Data de nascimento: {birth}")
    lines.append(f"- Sexo: {gender}")
    lines.append(f"- ID FHIR: {patient.get('id', '?')}")
    lines.append("")

    lines.append("## Diagnósticos")
    lines.append("")
    if conditions:
        for c in conditions:
            coding = c.get("code", {}).get("coding", [{}])[0]
            display = coding.get("display", "sem descrição")
            code = coding.get("code", "?")
            lines.append(f"- {display} ({code})")
    else:
        lines.append("- Nenhum diagnóstico registado.")
    lines.append("")

    lines.append("## Resultados laboratoriais")
    lines.append("")
    if observations:
        lines.append("| Data | Exame | Valor | LOINC |")
        lines.append("|------|-------|-------|-------|")
        sorted_obs = sorted(
            observations,
            key=lambda o: o.get("effectiveDateTime", o.get("issued", "")),
            reverse=True
        )
        for obs in sorted_obs[:20]:
            date = obs.get("effectiveDateTime", obs.get("issued", "?"))
            coding = obs.get("code", {}).get("coding", [{}])[0]
            display = coding.get("display", "?")
            loinc = coding.get("code", "?")
            vq = obs.get("valueQuantity", {})
            if vq:
                value = f"{vq.get('value', '?')} {vq.get('unit', '')}"
            elif obs.get("valueString"):
                value = obs["valueString"]
            else:
                value = "?"
            lines.append(f"| {date} | {display} | {value} | {loinc} |")
    else:
        lines.append("- Nenhuma observação registada.")
    lines.append("")

    lines.append("## Resumo clínico (gerado por LLM)")
    lines.append("")
    lines.append(summary)
    lines.append("")

    filename = f"resumo_clinico_{patient_id}.md"
    with open(filename, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"  Markdown: {filename}")
    return filename


def export_docx(patient_id, patient, conditions, observations, summary):
    """Exporta o resumo clínico como ficheiro Word (.docx)."""
    if not HAS_DOCX:
        print("  Word: ignorado (instalar com: pip install python-docx)")
        return None

    doc = Document()

    # Estilos base
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    # Título
    name_parts = patient.get("name", [{}])[0]
    given = " ".join(name_parts.get("given", []))
    family = name_parts.get("family", "")
    name = f"{given} {family}".strip() or "Desconhecido"

    title = doc.add_heading(f"Resumo Clínico: {name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadados
    meta_text = (
        f"Data de geração: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}  |  "
        f"Modelo: {MODEL}  |  "
        f"Servidor FHIR: {FHIR_BASE}"
    )
    meta_para = doc.add_paragraph(meta_text)
    meta_para.style.font.size = Pt(9)
    meta_para.style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Identificação
    doc.add_heading("Identificação do doente", level=1)
    birth = patient.get("birthDate", "desconhecida")
    gender = patient.get("gender", "desconhecido")
    pid = patient.get("id", "?")
    id_table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    id_table.autofit = True
    for row, (label, value) in enumerate([
        ("Nome", name),
        ("Data de nascimento", birth),
        ("Sexo", gender),
        ("ID FHIR", pid),
    ]):
        id_table.rows[row].cells[0].text = label
        id_table.rows[row].cells[1].text = value

    doc.add_paragraph("")

    # Diagnósticos
    doc.add_heading("Diagnósticos", level=1)
    if conditions:
        for c in conditions:
            coding = c.get("code", {}).get("coding", [{}])[0]
            display = coding.get("display", "sem descrição")
            code = coding.get("code", "?")
            doc.add_paragraph(f"{display} ({code})", style="List Bullet")
    else:
        doc.add_paragraph("Nenhum diagnóstico registado.")

    # Resultados
    doc.add_heading("Resultados laboratoriais", level=1)
    if observations:
        sorted_obs = sorted(
            observations,
            key=lambda o: o.get("effectiveDateTime", o.get("issued", "")),
            reverse=True
        )
        obs_table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
        obs_table.autofit = True
        hdr = obs_table.rows[0]
        for i, h in enumerate(["Data", "Exame", "Valor", "LOINC"]):
            hdr.cells[i].text = h

        for obs in sorted_obs[:20]:
            date = obs.get("effectiveDateTime", obs.get("issued", "?"))
            coding = obs.get("code", {}).get("coding", [{}])[0]
            display = coding.get("display", "?")
            loinc = coding.get("code", "?")
            vq = obs.get("valueQuantity", {})
            if vq:
                value = f"{vq.get('value', '?')} {vq.get('unit', '')}"
            elif obs.get("valueString"):
                value = obs["valueString"]
            else:
                value = "?"
            row = obs_table.add_row()
            row.cells[0].text = str(date)
            row.cells[1].text = str(display)
            row.cells[2].text = str(value)
            row.cells[3].text = str(loinc)
    else:
        doc.add_paragraph("Nenhuma observação registada.")

    doc.add_paragraph("")

    # Resumo LLM
    doc.add_heading("Resumo clínico (gerado por LLM)", level=1)
    for paragraph_text in summary.split("\n"):
        stripped = paragraph_text.strip()
        if stripped:
            doc.add_paragraph(stripped)

    # Aviso
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph(
        "AVISO: Este resumo foi gerado automaticamente por um modelo de linguagem (LLM) "
        "a partir de dados FHIR. Requer validação por um profissional de saúde antes de "
        "qualquer utilização clínica."
    )
    for run in disclaimer.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x33, 0x33)

    filename = f"resumo_clinico_{patient_id}.docx"
    doc.save(filename)
    print(f"  Word: {filename}")
    return filename


# ──────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────

def main():
    if len(sys.argv) < 2:
        print("Uso: python fhir_llm_summary.py <PATIENT_ID>")
        print("Exemplo: python fhir_llm_summary.py 131365376")
        sys.exit(1)

    patient_id = sys.argv[1]
    api_key = get_api_key()

    # 1. Obter dados FHIR
    entries = fetch_everything(patient_id)
    patient, conditions, observations, _ = extract_resources(entries)

    if not patient:
        print(f"ERRO: nao foi encontrado um recurso Patient para o ID {patient_id}.")
        sys.exit(1)

    # 2. Formatar contexto clinico
    clinical_context = build_clinical_context(patient, conditions, observations)
    print()
    print("=" * 60)
    print("DADOS CLINICOS EXTRAIDOS DO FHIR")
    print("=" * 60)
    print(clinical_context)
    print()

    # 3. Chamar o LLM
    summary = call_llm(clinical_context, api_key)
    print("=" * 60)
    print("RESUMO CLINICO GERADO PELO LLM")
    print("=" * 60)
    print(summary)
    print()

    # 4. Exportar ficheiros
    print("A exportar ficheiros:")
    export_json(patient_id, patient, conditions, observations, summary)
    export_markdown(patient_id, patient, conditions, observations, summary)
    export_docx(patient_id, patient, conditions, observations, summary)
    print()
    print("Concluído.")


if __name__ == "__main__":
    main()
